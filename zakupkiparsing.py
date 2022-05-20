# -*- coding: utf-8 -*-
# vim:fileencoding=utf-8
import datetime
import sqlite3
import docx
from bs4 import BeautifulSoup as BS
import requests
from fake_useragent import UserAgent

import re
import os
import time
import pandas as pd

import os, shutil

import os
import io
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfpage import PDFPage
import sys



def send_telegram(text: str):
    token = "5183755894:AAFJWAW0Xf3Vw5X2RLe-tJvqOBI1OKbZmjE"
    url = "https://api.telegram.org/bot"
    channel_id = "-1001676423488"
    url += token
    method = url + "/sendMessage"

    r = requests.post(method, data={
         "chat_id": channel_id,
         "text": text
          })

    if r.status_code != 200:
        raise Exception("post_text error")



def extract_text_from_pdf(pdf_path):
    resource_manager = PDFResourceManager()
    fake_file_handle = io.StringIO()
    converter = TextConverter(resource_manager, fake_file_handle)
    page_interpreter = PDFPageInterpreter(resource_manager, converter)

    with open(pdf_path, 'rb') as fh:
        for page in PDFPage.get_pages(fh,
                                      caching=True,
                                      check_extractable=True):
            try:
                page_interpreter.process_page(page)
            except:
                pass

        text = fake_file_handle.getvalue()

    # close open handles
    converter.close()
    fake_file_handle.close()

    if text:
        return text





link_main_pattern = 'http://zakupki.gov.ru'
# link_pattern_fz44_order_info = 'http://zakupki.gov.ru/epz/order/notice/ea44/view/common-info.html?regNumber='
link_pattern_fz44_result_info = 'http://zakupki.gov.ru/epz/order/notice/ea44/view/supplier-results.html?regNumber='
# link_pattern_fz44_documents = 'http://zakupki.gov.ru/epz/order/notice/ea44/view/documents.html?regNumber='
# link_pattern_fz44_contract_info = 'http://zakupki.gov.ru/epz/contract/extendedsearch/results.html?searchString=&orderNumber='
link_pattern_fz44_contruct = 'http://zakupki.gov.ru/epz/contract/contractCard/document-info.html?reestrNumber='
link_pattern_fz223 = 'http://zakupki.gov.ru/223/purchase/public/purchase/info/common-info.html?regNumber='
link_fz223_documents = 'http://zakupki.gov.ru/223/purchase/public/purchase/info/documents.html?regNumber='




def get_checkedkeys():
    f = open("chekced.txt","r")
    text=" "
    CheckedKeys=[]
    while text!="":
        text=f.readline()
        CheckedKeys.append(text.strip())
    f.close()
    return CheckedKeys


def get_page(link):
    page = requests.get(link, headers={'User-Agent': UserAgent().chrome})
    page.encoding = 'utf8'
    page = page.text
    soup = BS(page,"html.parser")
    while soup.title.text == '67756':
        time.sleep(5)
        get_page(link)
    return soup
def Check1(ID):
    f = open("ctru.txt", "r")
    line = f.readline()
    Keys = line.split("|")
    f.close()

    soup = get_page('https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=' + ID)
    divs_blocinfo1 = soup.find_all('tr', {"class": "tableBlock__row"})
    col = 1
    for i in divs_blocinfo1:
        conn = sqlite3.connect('IDGos.db')
        database = conn.cursor()
        soup=BS(str(i),"html.parser")
        divs_blocinfo2 = soup.find_all('td', {"class": "tableBlock__col"})
        try:

            for h in Keys:
                if h in str(divs_blocinfo2[1]):
                    try:
                        try:
                            database.execute("UPDATE IDSGOS SET 'Товар " + str(col) + "' = '" + h + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=" + ID + "'")
                        except:
                            pass
                        try:
                            database.execute("UPDATE IDSGOS SET 'Инфа " + str(col) + "' = '" + str(divs_blocinfo2[2].get_text()).strip().replace("    ","").replace("\n","") + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=" + ID + "'")
                        except:
                            pass
                        try:
                            database.execute("UPDATE IDSGOS SET 'Ед.Измерения " + str(col) + "' = '" + str(divs_blocinfo2[3].get_text()).strip() + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=" + ID + "'")
                        except:
                            pass
                        try:
                            database.execute("UPDATE IDSGOS SET 'Кол-во " + str(col) + "' = '" + divs_blocinfo2[4].get_text().strip() + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=" + ID + "'")
                        except:
                            pass
                        try:
                            database.execute("UPDATE IDSGOS SET 'Цена за ед. " + str(col) + "' = '" + divs_blocinfo2[5].get_text().strip() + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=" + ID + "'")
                        except:
                            pass
                        try:
                            database.execute("UPDATE IDSGOS SET 'Стоимость " + str(col) + "' = '" + divs_blocinfo2[6].get_text().strip() + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=" + ID + "'")
                        except:
                            pass
                        col += 1
                    except:
                        pass
            conn.commit()
        except:
            pass
def get_Infobyid(ID):

    conn = sqlite3.connect('IDGos.db')
    database = conn.cursor()
    database.execute("INSERT INTO IDSGOS (idgos, Дата) VALUES ('"+"https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber="+ID+"', '"+get_DatebyID(ID).strip()+"')")

    soup = get_page('https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber='+ID)

    divs_blocinfo = soup.find_all('div', {"class": "row blockInfo"})
    textinfo=''
    try:
        for i in divs_blocinfo:
            block=BS(str(i),"html.parser")
            checknameblock=block.find('h2', {"class": "blockInfo__title"})
            if(checknameblock.get_text() == 'Общая информация о закупке' or checknameblock.get_text() == 'Контактная информация' or checknameblock.get_text() == 'Начальная (максимальная) цена контракта'):
                blockinfoALL=block.find_all('section',{"class": "blockInfo__section section"})
                textinfo+=checknameblock.get_text()+": \n"
                for b in blockinfoALL:
                    blockone = BS(str(b), "html.parser")
                    first=blockone.find('span',{"class": "section__title"})
                    second=blockone.find('span',{"class": "section__info"})
                    try:
                        database.execute("ALTER TABLE IDSGOS ADD COLUMN '"+first.get_text().strip()+"' TEXT")
                    except:
                        #print("okey info")
                        pass
                    database.execute("UPDATE IDSGOS SET '"+first.get_text().strip()+"' = '"+second.get_text().strip()+"' WHERE idgos = '"+"https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber="+ID+"'")

                textinfo+="-------------------------------------------\n"
            elif(checknameblock.get_text() == 'Информация о процедуре закупки'):
                blockinfoALL = block.find_all('section', {"class": "blockInfo__section"})
                textinfo += checknameblock.get_text() + ": \n"
                for b in blockinfoALL:
                    blockone = BS(str(b), "html.parser")
                    first = blockone.find('span', {"class": "section__title"})
                    second = blockone.find('span', {"class": "section__info"})
                    try:
                        database.execute("ALTER TABLE IDSGOS ADD COLUMN '" + first.get_text().strip() + "' TEXT")
                    except:
                        #print("okey info1")
                        pass
                    database.execute("UPDATE IDSGOS SET '" + first.get_text().strip() + "' = '" + second.get_text().strip() + "' WHERE idgos = '" + "https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber="+ID + "'")
                textinfo += "-------------------------------------------\n"
    except:
        print("Error "+ID)
    conn.commit()
    Check1(ID)
    return textinfo






#Check1("0338300006822000029")
#Check1("0167200003422002025")
#print("==========================")
def get_allid(datestatus):
    date=''
    dateto=''
    if(datestatus==0):
        A=datetime.datetime.today()
        date=A.strftime("%d.%m.%Y")
        dateto =A.strftime("%d.%m.%Y")
    elif(datestatus==1):
        today = datetime.date.today()
        week_ago = today - datetime.timedelta(days=7)
        date=week_ago.strftime("%d.%m.%Y")
        A = datetime.datetime.today()
        dateto = A.strftime("%d.%m.%Y")

    f = open("key1.txt","r",encoding='utf-8')
    keys=f.readline().rsplit("|")
    numbers=set()
    CheckedKeys = get_checkedkeys()
    for key in keys:
        #print('https://zakupki.gov.ru/epz/order/extendedsearch/rss.html?searchString='+key+'&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_200&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom='+date+'&publishDateTo='+dateto)
        #soup = get_page('https://zakupki.gov.ru/epz/order/extendedsearch/?searchString='+key+'&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_200&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom='+date+'&publishDateTo='+dateto)
        #soup =get_page('https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString='+key+'&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_200&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom='+date+'&publishDateTo='+dateto)
        soup = get_page(
            'https://zakupki.gov.ru/epz/order/extendedsearch/rss.html?searchString='+key+'&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_50&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom='+date+'&publishDateTo='+dateto)
        #print('https://zakupki.gov.ru/epz/order/extendedsearch/rss.html?searchString='+key+'&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_50&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom='+date+'&publishDateTo='+dateto)
        contruct_link_div = soup.find_all('title')
        for i in contruct_link_div:
            if ("№" in i.get_text()):
                parts = i.get_text()
                parts = parts.rsplit("№")
                if (not parts[1] in CheckedKeys and len(parts[1]) == 19):
                    numbers.add(parts[1])
        #print(numbers)

    return numbers
def get_allid1(datestatus):
    date=''
    dateto=''
    if(datestatus==0):
        A=datetime.datetime.today()
        date=A.strftime("%d.%m.%Y")
        dateto =A.strftime("%d.%m.%Y")
    elif(datestatus==1):
        today = datetime.date.today()
        week_ago = today - datetime.timedelta(days=7)
        date=week_ago.strftime("%d.%m.%Y")
        A = datetime.datetime.today()
        dateto = A.strftime("%d.%m.%Y")

    f = open("key1.txt","r",encoding='utf-8')
    keys=f.readline().rsplit("|")
    numbers=set()
    CheckedKeys = get_checkedkeys()
    for key in keys:
        soup = get_page(
            'https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString='+key+'&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_200&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom='+date+'&publishDateTo='+dateto)
        contruct_link_div = soup.find_all('a', href=True)
        for i in contruct_link_div:
            if ("documents.html?regNumber=" in i.get('href')):
                parts = i.get('href')
                parts = parts.rsplit("=")
                if (not parts[1] in CheckedKeys and not "991111111222" in parts[1]):
                    numbers.add(parts[1])
        #print(numbers)

    return numbers

def get_filesbyid(Id):
    soup = get_page('https://zakupki.gov.ru/epz/order/notice/ea20/view/documents.html?regNumber='+Id)
    #document_link_div = soup.find_all('a', href=True)
    mydivs = soup.find_all("div", {"class": "blockFilesTabDocs"})
    soup1= BS(mydivs[0].__str__(),"html.parser")
    document_link_div = soup1.find_all('a', href=True)
    URLS = set()
    for i in document_link_div:
        if("filestore/public" in i.get('href')):
            URLS.add(i.get('href'))

    return URLS

def download_andcheck(URLS):
    b=0
    directory = 'files/'
    for i in URLS:
        link_to_file = i

        file = requests.get(i, stream=True, headers={'User-Agent': UserAgent().chrome})
        file.encoding = 'utf-8'

        if file.headers.get('Content-Disposition') is not None:
            file_extension = re.findall(r'\.[a-z][^\n]+', file.headers.get('Content-Disposition'))
        else:
            print(i)
            file_extension = ''
        parts = i.rsplit("=")
        #print(file_extension)
        if(len(file_extension)==0):
            break
        fileshare=file_extension[0].split(";")

        file_name = str(b) + fileshare[0][:-1]
        save_here = directory + '' + file_name
        if file.status_code == 200 and fileshare[0][:-1].find('.rar')==-1 and fileshare[0][:-1].find('.zip')==-1:
            with open(save_here, 'wb') as f:
                file.raw.decode_content = True
                shutil.copyfileobj(file.raw, f)
        print("Скачка завершена",save_here)

        if (fileshare[0][:-1] == ".xlsx"):
            try:
                pd.read_excel(save_here, engine='openpyxl')
            except ValueError as e:
                print("exception:", e)
            else:
                if(Check_xlsx(save_here)):
                    return True
                    continue
        if (fileshare[0][:-1] == ".xls"):
            try:
                pd.read_excel(save_here)
            except ValueError as e:
                print("exception:", e)
            else:
                if(Check_xls(save_here)):
                    return True
                    continue

        if (fileshare[0][:-1] == ".pdf"):
            try:
                text = extract_text_from_pdf(save_here)
            except ValueError as e:
                print("exception:", e)
            else:
                if(Check_pdf(save_here)):
                    return True
                    continue


        if(fileshare[0][:-1]==".docx" or fileshare[0][:-1]==".doc"):
            #if (file_extension[0][:-1] == '.doc'):
                #save_as_docx('C:\\Users\\Driver\\PycharmProjects\\pythonProject\\files\\'+file_name)
             #   fileshare[:-1] = '.docx'
            try:
                doc = docx.Document(save_here)
            except :
                print("not doc")
            else:
                if (Check_docx(save_here)):
                    return True
                    pass


        b += 1
def download_andcheckByOneId(URLS):
    b=0
    directory = 'files/'
    for i in URLS:
        link_to_file = i
        file = requests.get(i, stream=True, headers={'User-Agent': UserAgent().chrome})
        file.encoding = 'utf-8'

        if file.headers.get('Content-Disposition') is not None:
            file_extension = re.findall(r'\.[a-z][^\n]+', file.headers.get('Content-Disposition'))
        else:
            print(i)
            file_extension = ''
        parts = i.rsplit("=")
        #print(file_extension)
        if(len(file_extension)==0):
            break
        fileshare=file_extension[0].split(";")

        file_name = str(b) + fileshare[0][:-1]
        save_here = directory + '' + file_name
        if file.status_code == 200 and fileshare[0][:-1].find('.rar')==-1 and fileshare[0][:-1].find('.zip')==-1:
            with open(save_here, 'wb') as f:
                file.raw.decode_content = True
                shutil.copyfileobj(file.raw, f)

        if (fileshare[0][:-1] == ".xls" or fileshare[0][:-1] == ".xlsx"):
            try:
                text =pd.read_excel(save_here,engine='openpyxl')
            except ValueError as e:
                print("exception:", e)
            else:
                if(Check_xlsx(save_here)):
                    return True
                    continue

        if (fileshare[0][:-1] == ".pdf"):
            try:
                text = extract_text_from_pdf(save_here)
            except ValueError as e:
                print("exception:", e)
            else:
                if(Check_pdf(save_here)):
                    return True
                    continue


        if(fileshare[0][:-1]==".docx" or fileshare[0][:-1]==".doc"):
            #if (file_extension[0][:-1] == '.doc'):
                #save_as_docx('C:\\Users\\Driver\\PycharmProjects\\pythonProject\\files\\'+file_name)
             #   fileshare[:-1] = '.docx'
            try:
                doc = docx.Document(save_here)
            except ValueError as e:
                #print("exception:", e)
                pass
            else:
                if (Check_docx(save_here)):
                    return True
                    continue


        b += 1


def Check_docx(filename):
    f = open("key2.txt", "r", encoding='utf-8')
    keys = f.readline().rsplit("|")

    f = open("key3.txt", "r", encoding='utf-8')
    keys2 = f.readline().rsplit("|")

    doc = docx.Document(filename)
    text = ''
    for paragraph in doc.paragraphs:
        text+=paragraph.text
    try:
        all_tables = doc.tables
        for i, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    text += cell.text
    except:
        pass


    text=text.lower()
    print(len(text))
    #text = textract.process(filename)
    #print(text)

    for i in keys2:
        if ((i in text) or (i.lower() in text)):
            print("--------------" + i)
            return False

    return True
    # for i in keys:
    #     if((i in text) or (i.lower() in text)):
    #         print(i)
    #         return True

def Check_pdf(save_here):
    f = open("key2.txt", "r", encoding='utf-8')
    keys = f.readline().rsplit("|")
    f = open("key3.txt", "r", encoding='utf-8')
    keys2 = f.readline().rsplit("|")
    text = extract_text_from_pdf(save_here)
    text = str(text).lower()
    print(len(text))
    for i in keys2:
        if ((i in text) or (i.lower() in text)):
            print("--------------" + i)
            return False

    return True
    # for i in keys:
    #     if ((i in text) or (i.lower() in text)):
    #         return True

def Check_xlsx(save_here):
    f = open("key2.txt", "r", encoding='utf-8')
    keys = f.readline().rsplit("|")
    f = open("key3.txt", "r", encoding='utf-8')
    keys2 = f.readline().rsplit("|")
    xl = pd.ExcelFile(save_here)
    res = len(xl.sheet_names)
    xl.close()
    text=''
    for i in range(0,res):
        #text+=str(pd.read_excel('files/1.xlsx', i)).lower()
        excel_data =pd.read_excel(save_here, i,engine='openpyxl')
        check=excel_data.columns.tolist()
        for colum in check:
            vals=excel_data[colum].values.tolist()
            for val in vals:
                text+=str(val).lower()
        #print(excel_data[check[0]].values.tolist())
    text=text.lower()
    print(len(text))
    for i in keys2:
        if ((i in text) or (i.lower() in text)):
            print("--------------" + i)
            return False

    return True

    # for i in keys:
    #     if ((i in text) or (i.lower() in text)):
    #         #print(i,"    КЛЮЧxslx")
    #         return True

def Check_xls(save_here):
    f = open("key2.txt", "r", encoding='utf-8')
    keys = f.readline().rsplit("|")
    f = open("key3.txt", "r", encoding='utf-8')
    keys2 = f.readline().rsplit("|")
    xl = pd.ExcelFile(save_here)
    res = len(xl.sheet_names)
    xl.close()
    text=''
    for i in range(0,res):
        #text+=str(pd.read_excel('files/1.xlsx', i)).lower()
        excel_data =pd.read_excel(save_here, i)
        check=excel_data.columns.tolist()
        for colum in check:
            vals=excel_data[colum].values.tolist()
            for val in vals:
                text+=str(val).lower()
        #print(excel_data[check[0]].values.tolist())
    text=text.lower()
    print(len(text))
    for i in keys2:
        if ((i in text) or (i.lower() in text)):
            print("--------------"+i)
            return False

    return True
    # for i in keys:
    #     if ((i in text) or (i.lower() in text)):
    #         #print(i,"    КЛЮЧxslx")
    #         return True


#URLS = get_filesbyid('0351200005622000066')
#Check_docx('files/3.docx')
#download_andcheck(URLS)
def TodayGet(key1,chatId):
    print("1111")
    folder = 'C:\\Users\\Driver\\PycharmProjects\\pythonProject\\files'
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print('Failed to delete %s. Reason: %s' % (file_path, e))
    print(key1)

    numbers=[]
    A = datetime.datetime.today()
    date = A.strftime("%d.%m.%Y")
    dateto = A.strftime("%d.%m.%Y")
    date='11.04.2022'
    dateto='11.04.2022'
    soup = get_page(
        'https://zakupki.gov.ru/epz/order/extendedsearch/results.html?searchString=' + key1 + '&morphology=on&search-filter=Дате+размещения&pageNumber=1&sortDirection=false&recordsPerPage=_200&showLotsInfoHidden=false&sortBy=UPDATE_DATE&fz44=on&fz223=on&af=on&ca=on&pc=on&pa=on&currencyIdGeneral=-1&publishDateFrom=' + date + '&publishDateTo=' + dateto)
    contruct_link_div = soup.find_all('a', href=True)
    for i in contruct_link_div:
        if ("documents.html?regNumber=" in i.get('href')):
            parts = i.get('href')
            parts = parts.rsplit("=")
            if (not parts[1] in numbers):
                numbers.append(parts[1])
    print(numbers)
    for i in numbers:
        folder = 'C:\\Users\\Driver\\PycharmProjects\\pythonProject\\files'
        URLS = get_filesbyid(i)
        if (download_andcheck(URLS)):
            textinfofin = 'https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=' + i + '\n\n' + get_Infobyid(
                i)
            bot.send_message(chatId,textinfofin)
            break
            print(i)

        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))
        f = open("chekced.txt", "a",encoding='utf-8')
        f.write(i + "\n")
        f.close()
    bot.send_message(chatId, "Конец поиска")
def get_DatebyID(ID):
    soup = get_page('https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber=' + ID)
    divs_blocinfo = soup.find_all('div', {"class": "cardMainInfo__section col-6"})
    textinfo = ''
    for i in divs_blocinfo:
        block = BS(str(i), "html.parser")
        checknameblock = block.find('span', {"class": "cardMainInfo__title"})
        checkcontecnt = block.find('span', {"class": "cardMainInfo__content"})
        if(checknameblock.get_text()=='Размещено'):
            return checkcontecnt.get_text()





def send_telegram(text: str,id):
    token = "5183755894:AAFJWAW0Xf3Vw5X2RLe-tJvqOBI1OKbZmjE"
    url = "https://api.telegram.org/bot"
    url += token
    method = url + "/sendMessage"

    r = requests.post(method, data={
         "chat_id": str(id),
         "text": text
          })

    if r.status_code != 200:
        raise Exception("post_text error")

#ids=get_filesbyid('0817200000322004074')

#download_andcheck(ids)

def Checkdenied(URLS):
    b = 0
    directory = 'files/'
    for i in URLS:
        link_to_file = i

        file = requests.get(i, stream=True, headers={'User-Agent': UserAgent().chrome})
        file.encoding = 'utf-8'

        if file.headers.get('Content-Disposition') is not None:
            file_extension = re.findall(r'\.[a-z][^\n]+', file.headers.get('Content-Disposition'))
        else:
            print(i)
            file_extension = ''
        parts = i.rsplit("=")
        # print(file_extension)
        if (len(file_extension) == 0):
            break
        fileshare = file_extension[0].split(";")

        file_name = str(b) + fileshare[0][:-1]
        save_here = directory + '' + file_name
        if file.status_code == 200 and fileshare[0][:-1].find('.rar') == -1 and fileshare[0][:-1].find('.zip') == -1:
            with open(save_here, 'wb') as f:
                file.raw.decode_content = True
                shutil.copyfileobj(file.raw, f)
        print("Скачка завершена", save_here)
        if (fileshare[0][:-1] == ".xlsx"):
            try:
                pd.read_excel(save_here, engine='openpyxl')
            except ValueError as e:
                print("exception:", e)
            else:
                if (Check_xlsx(save_here)==False):
                    return False
                    continue
        if (fileshare[0][:-1] == ".xls"):
            try:
                pd.read_excel(save_here)
            except ValueError as e:
                print("exception:", e)
            else:
                if (Check_xls(save_here)==False):
                    return False
                    continue

        if (fileshare[0][:-1] == ".pdf"):
            try:
                text = extract_text_from_pdf(save_here)
            except ValueError as e:
                print("exception:", e)
            else:
                if (Check_pdf(save_here)==False):
                    return False
                    continue
        if (fileshare[0][:-1] == ".docx" or fileshare[0][:-1] == ".doc"):
            # if (file_extension[0][:-1] == '.doc'):
            # save_as_docx('C:\\Users\\Driver\\PycharmProjects\\pythonProject\\files\\'+file_name)
            #   fileshare[:-1] = '.docx'
            try:
                time.sleep(1)
                docx.Document(save_here)
            except:
                pass
            else:
                if (Check_docx(save_here)==False):
                    return False
                    continue
        b += 1
    return True


def Checkctru(id):
    page = requests.get("https://zakupki.gov.ru/epz/order/notice/ea20/view/common-info.html?regNumber="+id, headers={'User-Agent': UserAgent().chrome})
    page.encoding = 'utf8'
    page = page.text
    f = open("ctru.txt", "r")
    line=f.readline()
    Keys = line.split("|")
    f.close()
    page=page.strip()
    for i in Keys:
        if i in page:
            print(id,i)
            return True
    return False


if len(sys.argv) > 1:
        userid = sys.argv[1]

#print(Checkctru("0860200000822002219"))




Urls=get_filesbyid("0163200000322002142")
print(Checkdenied(Urls))

try:
    f = open('parsingstatus.txt','w')
    f.write('true')
    f.close()
    #get_Infobyid('0318200028122000359')
    #generate_xlsx("check2","08.04.2022")
    send_telegram("Парсинг запущен",userid)
    IDS=get_allid1(1)
    print(len(IDS))
    #print(IDS)
    for i in IDS:
        folder='.\\files'

        # if(download_andcheck(URLS)):
        #     get_Infobyid(i)
        #     print(i)

        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print('Failed to delete %s. Reason: %s' % (file_path, e))

        if (Checkctru(i)):
            URLS = get_filesbyid(i)
            if (Checkdenied(URLS)):
                print("+++++"+i+"+++++++")
                get_Infobyid(i)
            #print(i)
        f = open("chekced.txt", "a")
        f.write(i+"\n")
        f.close()
    send_telegram("Парсинг окончен",userid)
    f = open('parsingstatus.txt','w')
    f.write('false')
    f.close()
    print("Конец парсинга")
except Exception as exc:
    print(exc)
    f = open('parsingstatus.txt', 'w')
    f.write('false')
    f.close()
    send_telegram("При парсинге произошла ошибка, ошибка переданна разработчику", userid)
